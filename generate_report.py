"""
generate_report.py — Generate the project summary Word document.
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import os

doc = Document()

# ─── Style Configuration ─────────────────────────────────────────────
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)

for i in range(1, 5):
    h = doc.styles[f"Heading {i}"]
    h.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)  # dark blue
    h.font.name = "Calibri"

# Helper to add styled paragraphs
def add_para(text, bold=False, italic=False, size=None, alignment=None, space_after=Pt(6)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = size
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_after = space_after
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Cm(1.27 + level * 1.27)
    return p

def add_code_block(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    return p


# ═══════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════════════
doc.add_paragraph()  # spacer
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = p.add_run("Collibra \"Ask Your Data\" Challenge")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

p = doc.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = p.add_run("Project Guide & Technical Documentation")
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = p.add_run("Next Level Challenge Hackathon")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = p.add_run("A step-by-step guide to building a governed, deterministic\ndata question-answering prototype using Collibra context")
run.italic = True
run.font.size = Pt(12)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading("Table of Contents", level=1)
toc_items = [
    "1. Executive Summary",
    "2. Problem Statement: Why This Challenge Exists",
    "3. Architecture Overview: The Four-Layer Model",
    "4. Step-by-Step Implementation Guide",
    "    4.1 Environment Setup",
    "    4.2 Cloning the Repository & Installing Dependencies",
    "    4.3 Configuring API Connections",
    "    4.4 Ingesting Governed Context from Collibra",
    "    4.5 Building the RAG Pipeline (LLM-based)",
    "    4.6 Running the Before/After Demo",
    "    4.7 Implementing the Rulebook",
    "    4.8 Building the Deterministic Engine",
    "    4.9 Verification: 20-Run Determinism Test",
    "    4.10 Overlap Analysis: Proving No Double-Counting",
    "5. Why Context Matters: Before vs. After",
    "6. The Rulebook: Why and How",
    "7. Results Summary",
    "8. How to Replicate This Project",
    "9. Appendix: File Reference",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    if item.startswith("    "):
        p.paragraph_format.left_indent = Cm(1.27)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading("1. Executive Summary", level=1)

add_para(
    "This document describes the end-to-end implementation of a prototype built for the "
    "Collibra \"Ask Your Data\" hackathon challenge. The goal was to build a system that "
    "answers plain-language business questions about Collibra's preview program data, "
    "using governed context from the Collibra Data Intelligence Platform to ensure accuracy, "
    "consistency, and trustworthiness.",
    space_after=Pt(8)
)

add_para(
    "The core insight is simple but powerful: when a business user asks \"How many active testers "
    "do we have in ongoing programs?\", an AI system without governed context will guess what "
    "\"active tester\" and \"ongoing\" mean. It will produce plausible-looking but wrong answers. "
    "Our prototype uses Collibra's governed business definitions to ensure the system "
    "interprets the question exactly as the organization intends.",
    space_after=Pt(8)
)

doc.add_heading("What We Built", level=2)
add_bullet("A RAG (Retrieval-Augmented Generation) pipeline that injects Collibra-governed "
           "context into LLM prompts, so the LLM generates correct SQL grounded in official definitions.")
add_bullet("A deterministic rulebook engine that translates Collibra's business definitions "
           "into executable Python code, producing identical answers on every run — no LLM required.")
add_bullet("A before/after demonstration showing that without context, the LLM returns 0 results "
           "(wrong), while with context, it returns 303 active testers (correct).")

doc.add_heading("Key Result", level=2)
add_para(
    "303 active testers in ongoing programs. This number was verified across 20 consecutive runs "
    "with zero variance. The answer is derived directly from Collibra's governed definition of "
    "\"Active Tester\" (business term asset ID: 0198c234-11fe-73ff-be9b-c91312850031) and "
    "\"Active Tester Flag\" (measure asset ID: 019c9f78-2633-7377-9050-c1b3d84eb68d).",
    space_after=Pt(8)
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 2. PROBLEM STATEMENT
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading("2. Problem Statement: Why This Challenge Exists", level=1)

add_para(
    "Modern organizations sit on vast amounts of data, but business users struggle to extract "
    "value from it. They ask questions in natural language — \"How many active testers do we have?\" "
    "— but the underlying data lives in databases with cryptic column names like Z_ACT_CMP, "
    "Z_SRV_CMP, and Z_PRJ_STAT. Without knowing what these columns mean, any automated system "
    "will guess — and likely guess wrong.",
    space_after=Pt(8)
)

doc.add_heading("The Gap", level=2)
add_para(
    "Large Language Models (LLMs) are excellent at generating SQL from natural language, but they "
    "have a critical weakness: they do not know the organization's specific business vocabulary. "
    "When an LLM sees a column called Z_PRJ_STAT, it might guess that 'ACTIVE' is a valid value — "
    "but in this dataset, the actual values are 'Ongoing' and 'Finished'. This is not a limitation "
    "of the model's intelligence; it is a fundamental information gap.",
    space_after=Pt(8)
)

doc.add_heading("Collibra's Role", level=2)
add_para(
    "Collibra is a Data Intelligence Platform that serves as the single source of truth for "
    "business vocabulary, data definitions, and metadata governance. In Collibra, every business "
    "term has an official definition, every metric has a calculation formula, and every physical "
    "database column is mapped to its business meaning through a traceable lineage.",
    space_after=Pt(8)
)

add_para(
    "The challenge asks: can we bridge the gap between human questions and machine data by using "
    "Collibra's governed context as the \"Rosetta Stone\" that translates business language into "
    "precise data operations?",
    space_after=Pt(8)
)

doc.add_heading("Why It Matters", level=2)
add_bullet("Without context: The LLM guesses 'ACTIVE' as a project status and returns 0 results.")
add_bullet("With context: The system knows 'ongoing' means Z_PRJ_STAT = 'Ongoing' (20 projects, "
           "not archived), and 'active tester' has a precise three-criteria definition. It returns "
           "303 — the correct answer.")
add_bullet("This is not an edge case. Every organization has its own vocabulary. Without governance, "
           "every AI-powered data tool is essentially hallucinating business logic.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 3. ARCHITECTURE OVERVIEW
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading("3. Architecture Overview: The Four-Layer Model", level=1)

add_para(
    "The solution follows Collibra's four-layer data governance architecture. Understanding these "
    "layers is essential because each layer adds a different type of context that makes the final "
    "answer trustworthy.",
    space_after=Pt(8)
)

doc.add_heading("Layer 1: Business Vocabulary (Glossary + Metrics)", level=2)
add_para(
    "This is the \"what\" layer. It contains governed definitions of business terms and metrics. "
    "For example, the term \"Active Tester\" is defined as: \"An active participant in a preview "
    "program who meets at least one of three criteria: (1) submitted 3 or more feedback tickets, "
    "(2) completed 2 or more surveys, or (3) completed more than half of their assigned activities.\"",
    space_after=Pt(6)
)
add_para(
    "Without this layer, the concept of \"active tester\" is ambiguous — it could mean anything "
    "from \"logged in recently\" to \"submitted at least one ticket.\"",
    space_after=Pt(8)
)

doc.add_heading("Layer 2: Logical Layer (Data Entities & Attributes)", level=2)
add_para(
    "This layer maps business concepts to logical data structures. For instance, the business term "
    "\"Completed Activities\" maps to a logical attribute that represents the count of activities "
    "a participant has finished. This layer acts as a bridge, ensuring that business terms are "
    "connected to data concepts in a technology-agnostic way.",
    space_after=Pt(8)
)

doc.add_heading("Layer 3: Physical Layer (Schemas, Tables, Columns)", level=2)
add_para(
    "This layer describes the actual database structure — table names, column names, and their "
    "descriptions. For example, physical column zcc_prt_mtrc.Z_ACT_CMP maps to the logical "
    "attribute \"Completed Activities\", which in turn maps to the business term of the same name. "
    "The column name Z_ACT_CMP is meaningless on its own; the physical layer gives it a human-"
    "readable description: \"Number of activities the participant has completed.\"",
    space_after=Pt(8)
)

doc.add_heading("Layer 4: Data Layer (Databricks via Delta Sharing)", level=2)
add_para(
    "This is the actual data — 8 tables hosted in Databricks and accessed via Delta Sharing, "
    "a secure, open protocol for sharing data across platforms without copying it. The tables "
    "contain preview program data: participant metrics (zcc_prt_mtrc), feedback tickets "
    "(zcc_tkt_itm), project headers (zcc_prj_hdr), activity status (zcc_act_stat), and more.",
    space_after=Pt(8)
)

doc.add_heading("How the Layers Connect", level=2)
add_para(
    "The power of this architecture is the chain of traceability:", space_after=Pt(4)
)
add_para(
    "Business Question -> Business Term (Layer 1) -> Logical Attribute (Layer 2) -> "
    "Physical Column (Layer 3) -> Actual Data (Layer 4)",
    bold=True, space_after=Pt(6)
)
add_para(
    "When a user asks about \"active testers\", the system traces this chain: \"Active Tester\" "
    "(business term) -> \"Active Tester Flag\" (measure with criteria) -> Z_ACT_CMP, Z_ACT_INC, "
    "Z_ACT_BLK, Z_ACT_OPT, Z_SRV_CMP, Z_TKT_NR (physical columns) -> actual data in Databricks. "
    "Every step is governed, auditable, and traceable.",
    space_after=Pt(8)
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 4. STEP-BY-STEP IMPLEMENTATION GUIDE
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading("4. Step-by-Step Implementation Guide", level=1)

add_para(
    "This section walks through every step required to replicate this project from scratch. "
    "Each step explains what is being done, why it is necessary, and what to verify before "
    "moving to the next step.",
    space_after=Pt(8)
)

# --- 4.1 ---
doc.add_heading("4.1 Environment Setup", level=2)
add_para("Purpose: Create an isolated Python environment so dependencies do not conflict with "
         "other projects on your machine.", space_after=Pt(6))
add_para("What you need:", bold=True)
add_bullet("Python 3.10 or higher (we used 3.12)")
add_bullet("Git (to clone the hackathon repository)")
add_bullet("A terminal (PowerShell on Windows, Terminal on macOS/Linux)")

add_para("Why a virtual environment?", bold=True, space_after=Pt(4))
add_para(
    "Python packages can conflict with each other. A virtual environment (venv) creates a "
    "sandboxed Python installation where you install only the packages this project needs. "
    "This prevents version conflicts and makes the project reproducible.",
    space_after=Pt(8)
)

# --- 4.2 ---
doc.add_heading("4.2 Cloning the Repository & Installing Dependencies", level=2)
add_para("Purpose: Get the starter code and reference documentation from the hackathon organizers.",
         space_after=Pt(6))

add_para("Commands:", bold=True)
add_code_block("git clone https://github.com/erkantrt/NextLevelChallenge.git\ncd NextLevelChallenge\npython -m venv .venv\n.venv\\Scripts\\activate   # Windows\npip install delta-sharing pandas requests anthropic beautifulsoup4 lxml python-dotenv pandasql")

add_para("What the repository contains:", bold=True)
add_bullet("README.md — overview of the challenge and architecture")
add_bullet("Collibra_API_Guide.md — reference for the Collibra REST API endpoints")
add_bullet("Delta_sharing_Guide.md — instructions for connecting to Databricks data")

add_para("Key dependencies and why we need each:", bold=True)
add_bullet("delta-sharing — reads data from Databricks tables via the Delta Sharing protocol")
add_bullet("pandas — data manipulation and analysis (the core computation library)")
add_bullet("requests — makes HTTP calls to the Collibra REST API")
add_bullet("anthropic — client library for the Claude LLM (used in the RAG pipeline)")
add_bullet("beautifulsoup4 + lxml — strips HTML tags from Collibra attribute values "
           "(Collibra stores some definitions as HTML)")
add_bullet("python-dotenv — loads credentials from a .env file so they are not hardcoded")
add_bullet("pandasql — lets us run SQL queries directly against pandas DataFrames")

doc.add_page_break()

# --- 4.3 ---
doc.add_heading("4.3 Configuring API Connections", level=2)
add_para("Purpose: Set up credentials for all three external systems the prototype connects to.",
         space_after=Pt(6))

add_para("Three connections are required:", bold=True)

doc.add_heading("A. Delta Sharing (Databricks Data)", level=3)
add_para(
    "Delta Sharing is an open protocol developed by Databricks for securely sharing data without "
    "copying it. A provider (the hackathon organizers) shares specific tables, and a recipient "
    "(us) receives a config.json file containing a bearer token and endpoint URL.",
    space_after=Pt(6)
)
add_para("config.json structure:", bold=True)
add_code_block('{\n  "shareCredentialsVersion": 1,\n  "bearerToken": "<your-token>",\n'
               '  "endpoint": "https://ohio.cloud.databricks.com/api/2.0/delta-sharing/"\n}')
add_para(
    "This file should be placed in the project root directory. It grants read-only access to "
    "8 tables in the centercode_share.centercode schema.",
    space_after=Pt(6)
)

doc.add_heading("B. Collibra REST API", level=3)
add_para(
    "The Collibra API provides programmatic access to all governed assets — business terms, "
    "measures, logical entities, physical tables, columns, and the relations between them. "
    "We use Basic Authentication (username + password) against the hackathon instance at "
    "https://next.collibra.com.",
    space_after=Pt(6)
)

doc.add_heading("C. LLM API (Anthropic Claude)", level=3)
add_para(
    "The RAG pipeline uses an LLM to generate SQL from natural language questions. We chose "
    "Anthropic's Claude (claude-sonnet-4-20250514) for its strong instruction-following and "
    "code generation capabilities. The API key is stored in the .env file.",
    space_after=Pt(6)
)

add_para(".env file structure:", bold=True)
add_code_block("COLLIBRA_USERNAME=<your-username>\nCOLLIBRA_PASSWORD=<your-password>\n"
               "ANTHROPIC_API_KEY=<your-api-key>\nDELTA_SHARING_PROFILE=config.json\nDELTA_SHARING_LIMIT=0")

add_para("Verification:", bold=True, space_after=Pt(4))
add_para(
    "Before proceeding, verify each connection independently. Test Delta Sharing by listing "
    "available tables. Test Collibra by fetching the communities endpoint. Test the LLM by "
    "sending a simple prompt. Each connection should return a successful response (HTTP 200). "
    "If any connection fails, the rest of the pipeline cannot function.",
    space_after=Pt(8)
)

doc.add_page_break()

# --- 4.4 ---
doc.add_heading("4.4 Ingesting Governed Context from Collibra", level=2)
add_para("Purpose: Extract all relevant business definitions, metrics, column descriptions, "
         "and relations from Collibra into a local JSON file that the prototype can use.",
         space_after=Pt(6))

add_para("File: collibra_client.py", bold=True, space_after=Pt(4))

add_para("What this script does:", bold=True)
add_bullet("Connects to the Collibra REST API using Basic Authentication")
add_bullet("Fetches assets from four domains: Glossary (business terms), Metrics (measures), "
           "Logical (data entities/attributes), and Physical (schemas, tables, columns)")
add_bullet("For each asset, retrieves its attributes (Definition, Description, Synonym, etc.)")
add_bullet("Retrieves relations between assets across all four layers")
add_bullet("Builds a column map that links every physical column to its business term")
add_bullet("Saves everything to governed_context.json for offline use")

add_para("Why this step is critical:", bold=True, space_after=Pt(4))
add_para(
    "This is where the \"magic\" happens. Without this step, we have data with meaningless column "
    "names and an LLM that can only guess. After this step, we have a complete dictionary that "
    "maps every cryptic column name (e.g., Z_ACT_CMP) to its official business meaning "
    "(\"Number of activities the participant has completed\").",
    space_after=Pt(6)
)

add_para("What we extracted:", bold=True)
add_bullet("52 business terms from the Glossary domain")
add_bullet("4 measures from the Metrics domain")
add_bullet("52 logical assets (data entities and attributes)")
add_bullet("74 physical assets (schemas, tables, columns)")
add_bullet("230 relations connecting assets across all four layers")
add_bullet("65 column-level mappings (physical column -> business term)")

add_para("Technical note — HTML stripping:", bold=True, space_after=Pt(4))
add_para(
    "Collibra stores some attribute values as HTML (e.g., definitions with bold text or lists). "
    "We use BeautifulSoup to strip HTML tags and extract plain text. We also encountered a bug "
    "where some numeric attribute values (e.g., 255.0) caused a TypeError because BeautifulSoup "
    "expects strings. The fix was to add a type check: if the value is not a string, convert it "
    "to a string before processing.",
    space_after=Pt(8)
)

add_para("Command to run:", bold=True)
add_code_block("python collibra_client.py")
add_para("Output: governed_context.json (used by all downstream components)", space_after=Pt(8))

doc.add_page_break()

# --- 4.5 ---
doc.add_heading("4.5 Building the RAG Pipeline (LLM-based)", level=2)
add_para("Purpose: Create the LLM-powered question-answering system that uses governed context "
         "to generate accurate SQL and answers.", space_after=Pt(6))

add_para("File: ask_your_data.py", bold=True, space_after=Pt(4))

add_para("How the RAG pipeline works:", bold=True)
add_bullet("Step 1: Load the governed context from governed_context.json")
add_bullet("Step 2: Load all 8 tables from Databricks via Delta Sharing into pandas DataFrames")
add_bullet("Step 3: Build a system prompt that includes all business term definitions, "
           "measure calculation rules, column mappings, and table schemas")
add_bullet("Step 4: Send the user's question along with this system prompt to Claude")
add_bullet("Step 5: Claude generates SQL grounded in the governed definitions")
add_bullet("Step 6: Execute the generated SQL against the pandas DataFrames using pandasql")
add_bullet("Step 7: Return the answer with full traceability (which definitions were used, "
           "which columns were mapped, which relations were traversed)")

add_para("What makes this a RAG system:", bold=True, space_after=Pt(4))
add_para(
    "RAG stands for Retrieval-Augmented Generation. Instead of relying solely on the LLM's "
    "training data (which does not know our specific business vocabulary), we retrieve the "
    "relevant governed context from Collibra and augment the LLM's prompt with it. The LLM "
    "then generates a response that is grounded in official definitions rather than guesses.",
    space_after=Pt(6)
)

add_para("Key design decisions:", bold=True)
add_bullet("Temperature = 0: We set the LLM's temperature to 0 to minimize randomness and "
           "produce more deterministic outputs.")
add_bullet("Explicit instructions: The system prompt tells the LLM to cite governed definitions, "
           "flag ungoverned terms, and always show its reasoning.")
add_bullet("SQL execution: We use pandasql (SQLite dialect) to execute the LLM-generated SQL "
           "directly against the pandas DataFrames, avoiding the need for a separate database.")

doc.add_page_break()

# --- 4.6 ---
doc.add_heading("4.6 Running the Before/After Demo", level=2)
add_para("Purpose: Demonstrate the value of governed context by comparing LLM answers with and "
         "without Collibra definitions.", space_after=Pt(6))

add_para("Question: \"How many active testers do we have in the ongoing programs?\"", 
         bold=True, italic=True, space_after=Pt(8))

doc.add_heading("WITHOUT Context (the \"Before\")", level=3)
add_para(
    "The LLM receives only the raw table schemas — column names and data types, but no "
    "definitions or business context. Here is what happened:",
    space_after=Pt(6)
)
add_bullet("The LLM guessed that Z_PRJ_STAT = 'ACTIVE' (wrong — the actual value is 'Ongoing')")
add_bullet("The LLM guessed that Z_ENG_ST = 'ACTIVE' (wrong — no such value exists)")
add_bullet("Execution returned 0 results because the filter matched nothing")
add_bullet("The LLM confidently stated: \"There are zero active testers\" — plausible but completely wrong")

add_para(
    "This is the core danger of using AI on enterprise data without governance: the answer looks "
    "right but is based on fabricated filter values.",
    italic=True, space_after=Pt(8)
)

doc.add_heading("WITH Context (the \"After\")", level=3)
add_para(
    "The LLM receives the full governed context: business term definitions, measure formulas, "
    "column descriptions, and table schemas. Here is what happened:",
    space_after=Pt(6)
)
add_bullet("The LLM used the governed definition of \"Active Tester\" — three specific criteria")
add_bullet("It used Z_PRJ_STAT = 'Ongoing' (the correct value from the governed definition)")
add_bullet("It applied the three Active Tester Flag criteria using the correct column names")
add_bullet("It correctly used OR logic (any one criterion met = active)")
add_bullet("Result: 303 active testers in ongoing programs (with the 5000-row-limited dataset)")

add_para("Why the difference matters:", bold=True, space_after=Pt(4))
add_para(
    "The same LLM, the same question, the same data — the only difference is context. This "
    "demonstrates that context is not optional; it is the difference between a useful tool and "
    "a dangerous one.",
    space_after=Pt(8)
)

doc.add_page_break()

# --- 4.7 ---
doc.add_heading("4.7 Implementing the Rulebook", level=2)
add_para("Purpose: Encode Collibra's governed definitions into deterministic, executable Python "
         "code that produces the same answer every time.", space_after=Pt(6))

add_para("File: rulebook.py", bold=True, space_after=Pt(4))

add_para("What the rulebook contains:", bold=True)
add_bullet("The verbatim Collibra definition of \"Active Tester\" (business term)")
add_bullet("The verbatim Collibra definition of \"Active Tester Flag\" (measure)")
add_bullet("The governed definition of \"Project Status\" values ('Ongoing' vs. 'Finished')")
add_bullet("A complete column map linking every physical column to its governed description")
add_bullet("The three criteria for the Active Tester Flag, with exact column names and thresholds")
add_bullet("All Collibra asset IDs for full traceability back to the source")

add_para("Why not just use the LLM?", bold=True, space_after=Pt(4))
add_para(
    "While the RAG pipeline with governed context produces correct answers, it still involves "
    "an LLM — which introduces a small but non-zero risk of variation. LLMs can rephrase "
    "queries slightly differently between runs, add extra filters, or interpret edge cases "
    "inconsistently. For a question where the business definition is fully specified (as it is "
    "for \"Active Tester\"), we can encode the logic directly in code and eliminate the LLM "
    "entirely for that specific question.",
    space_after=Pt(6)
)

add_para("The rulebook pattern:", bold=True, space_after=Pt(4))
add_para(
    "1. Read the governed definition from Collibra.\n"
    "2. Identify the criteria, thresholds, columns, and filter values.\n"
    "3. Verify them against actual data (e.g., confirm that Z_PRJ_STAT actually contains "
    "'Ongoing', not 'Active').\n"
    "4. Encode the verified logic as Python constants and code.\n"
    "5. Document every decision with the Collibra asset ID as the authoritative source.",
    space_after=Pt(6)
)

add_para(
    "The rulebook is not a replacement for the RAG pipeline — it is a complement. The RAG "
    "pipeline handles open-ended questions where the user might ask anything. The rulebook "
    "handles known, governed questions where the answer must be exact and repeatable.",
    italic=True, space_after=Pt(8)
)

doc.add_page_break()

# --- 4.8 ---
doc.add_heading("4.8 Building the Deterministic Engine", level=2)
add_para("Purpose: Implement the Active Tester computation as pure pandas code, using "
         "the rulebook definitions, with no LLM in the loop.", space_after=Pt(6))

add_para("File: deterministic_engine.py", bold=True, space_after=Pt(4))

add_para("How it works:", bold=True)
add_bullet("Step 1: Load three tables from Databricks — zcc_prt_mtrc (participant metrics), "
           "zcc_tkt_itm (feedback tickets), zcc_prj_hdr (project headers)")
add_bullet("Step 2: Filter to ongoing projects by joining with zcc_prj_hdr where "
           "Z_PRJ_STAT = 'Ongoing' (20 projects)")
add_bullet("Step 3: Apply Criterion 1 — count distinct ticket numbers (Z_TKT_NR) per "
           "participant (Z_SMTP_ADR) and select those with 3 or more")
add_bullet("Step 4: Apply Criterion 2 — sum completed surveys (Z_SRV_CMP) per participant "
           "and select those with 2 or more")
add_bullet("Step 5: Apply Criterion 3 — sum completed activities (Z_ACT_CMP) and compare against "
           "the sum of incomplete (Z_ACT_INC) + blocked (Z_ACT_BLK) + opted-out (Z_ACT_OPT) "
           "activities; select those where completed outnumber the rest")
add_bullet("Step 6: Combine using Python set union (|) — a participant is active if they meet "
           "ANY ONE of the three criteria")
add_bullet("Step 7: Return the count along with full provenance (criteria breakdown, governed "
           "sources, asset IDs, tables and columns used)")

add_para("Why set union prevents double-counting:", bold=True, space_after=Pt(4))
add_para(
    "A common concern with OR logic is: if Person A meets both Criterion 1 and Criterion 2, "
    "are they counted twice? The answer is no, because we use Python sets. Each criterion "
    "produces a set of email addresses (the unique participant identifier). The final answer "
    "is the union of all three sets. By definition, a set contains each element at most once, "
    "regardless of how many source sets it appeared in.",
    space_after=Pt(6)
)
add_para("Overlap analysis from our data:", bold=True)
add_bullet("Criterion 1 (tickets >= 3): 73 unique participants")
add_bullet("Criterion 2 (surveys >= 2): 284 unique participants")
add_bullet("Criterion 3 (completed > rest): 64 unique participants")
add_bullet("Naive sum (if we double-counted): 421")
add_bullet("C1 AND C2 overlap: 57 participants")
add_bullet("C1 AND C3 overlap: 28 participants")
add_bullet("C2 AND C3 overlap: 61 participants")
add_bullet("In all three criteria: 28 participants")
add_bullet("Set union (actual, no duplicates): 303")
add_bullet("Overlap removed: 118 participants (421 - 303)")

doc.add_page_break()

# --- 4.9 ---
doc.add_heading("4.9 Verification: 20-Run Determinism Test", level=2)
add_para("Purpose: Prove that the deterministic engine produces identical results on every run.",
         space_after=Pt(6))

add_para("File: test_deterministic.py", bold=True, space_after=Pt(4))

add_para("Methodology:", bold=True)
add_bullet("Run the deterministic engine 20 times consecutively")
add_bullet("Record the answer and criteria breakdown for each run")
add_bullet("Compare all 20 results for exact equality")
add_bullet("Save results to deterministic_test_results.json for audit")

add_para("Results:", bold=True)
add_bullet("All 20 runs returned exactly 303 active testers")
add_bullet("Criteria breakdown was identical across all runs: C1=73, C2=284, C3=64")
add_bullet("Zero variance — the engine is fully deterministic")

add_para("Why determinism matters:", bold=True, space_after=Pt(4))
add_para(
    "In data governance, reproducibility is non-negotiable. If two people ask the same question "
    "at the same time, they must get the same answer. If an auditor re-runs a report, the number "
    "must match. The deterministic engine guarantees this because it uses fixed logic (from the "
    "rulebook) applied to fixed data (from Databricks). There is no randomness anywhere in "
    "the pipeline.",
    space_after=Pt(8)
)

# --- 4.10 ---
doc.add_heading("4.10 Overlap Analysis: Proving No Double-Counting", level=2)
add_para("Purpose: Mathematically verify that the OR logic does not inflate the count by "
         "counting the same participant multiple times.", space_after=Pt(6))

add_para(
    "We ran a dedicated overlap analysis script that computed the intersection of every pair "
    "of criteria sets and the intersection of all three. The results (shown in Section 4.8) "
    "confirm that 118 participants appear in more than one criterion set, and the set union "
    "correctly counts each of them only once.",
    space_after=Pt(6)
)

add_para(
    "Mathematical proof: |A OR B OR C| = |A| + |B| + |C| - |A AND B| - |A AND C| - |B AND C| "
    "+ |A AND B AND C|. Substituting: 73 + 284 + 64 - 57 - 28 - 61 + 28 = 303. This matches "
    "our set union result exactly, confirming the computation is correct.",
    space_after=Pt(8)
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 5. WHY CONTEXT MATTERS
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading("5. Why Context Matters: Before vs. After", level=1)

add_para(
    "The single most important takeaway from this project is that governed context transforms "
    "AI-powered data tools from dangerous to trustworthy. Here is a concrete comparison:",
    space_after=Pt(8)
)

# Table
table = doc.add_table(rows=7, cols=3, style="Light Shading Accent 1")
headers = ["Aspect", "Without Context", "With Context"]
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h

rows_data = [
    ["Interpretation of\n\"active tester\"",
     "Guessed: anyone with\nZ_ENG_ST = 'ACTIVE'",
     "Governed: 3-criteria OR\nlogic from Collibra definition"],
    ["Interpretation of\n\"ongoing\"",
     "Guessed: Z_PRJ_STAT =\n'ACTIVE'",
     "Governed: Z_PRJ_STAT =\n'Ongoing' (not archived)"],
    ["SQL filter values",
     "Fabricated ('ACTIVE')",
     "Verified ('Ongoing',\n'Finished')"],
    ["Result", "0 testers (wrong)", "303 testers (correct)"],
    ["Traceability", "None — black box", "Full — every definition,\ncolumn, and relation cited"],
    ["Reproducibility", "LLM may vary between\nruns", "Deterministic — 20/20\nruns identical"],
]
for r_idx, row_data in enumerate(rows_data):
    for c_idx, cell_text in enumerate(row_data):
        table.rows[r_idx + 1].cells[c_idx].text = cell_text

doc.add_paragraph()  # spacer

add_para(
    "The context bridge is what makes the difference. Collibra's governed definitions act as "
    "constraints that prevent the AI from hallucinating business logic. Instead of \"I think "
    "active means this,\" the system says \"according to the governed definition (asset ID "
    "0198c234-11fe-73ff-be9b-c91312850031), active tester means this.\"",
    space_after=Pt(8)
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 6. THE RULEBOOK: WHY AND HOW
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading("6. The Rulebook: Why and How", level=1)

doc.add_heading("Why We Implemented a Rulebook", level=2)
add_para(
    "The RAG pipeline is powerful but has an inherent limitation: it relies on an LLM to "
    "interpret definitions and generate SQL. Even with governed context injected into the "
    "prompt, the LLM is making a translation — from English definitions to SQL code. "
    "This translation step introduces potential for subtle errors:",
    space_after=Pt(4)
)
add_bullet("The LLM might use >= instead of > for a threshold (off-by-one errors)")
add_bullet("The LLM might aggregate at the wrong grain (per-project vs. per-participant)")
add_bullet("The LLM might miss one of three criteria in a complex OR condition")
add_bullet("Different runs might produce slightly different SQL despite identical inputs")

add_para(
    "For governed metrics where the definition is fully specified, we can eliminate this risk "
    "entirely by encoding the logic directly. The rulebook is the bridge between Collibra's "
    "human-readable definitions and machine-executable code.",
    space_after=Pt(8)
)

doc.add_heading("How the Rulebook Works", level=2)

add_para("The rulebook follows a strict process:", bold=True, space_after=Pt(4))

add_para("1. Extract the governed definition from Collibra:", bold=True, space_after=Pt(4))
add_para(
    "We read the \"Active Tester\" business term (asset 0198c234-11fe-73ff-be9b-c91312850031) "
    "and \"Active Tester Flag\" measure (asset 019c9f78-2633-7377-9050-c1b3d84eb68d) from the "
    "Collibra API. These definitions specify exactly what \"active\" means.",
    space_after=Pt(6)
)

add_para("2. Identify the data elements:", bold=True, space_after=Pt(4))
add_para(
    "The Collibra relations graph shows that Active Tester Flag is linked to specific physical "
    "columns: Z_ACT_BLK (blocked activities), Z_ACT_INC (incomplete activities), Z_ACT_OPT "
    "(opted-out activities), Z_ACT_CMP (completed activities), Z_SRV_CMP (completed surveys), "
    "and ticket counts from zcc_tkt_itm. These relations were discovered by querying the "
    "Collibra relations API.",
    space_after=Pt(6)
)

add_para("3. Verify against actual data:", bold=True, space_after=Pt(4))
add_para(
    "We ran discovery scripts (discover.py, discover2.py) that loaded actual data from "
    "Databricks to verify: (a) the column names exist and are spelled correctly, (b) the "
    "filter values match reality (Z_PRJ_STAT contains 'Ongoing' and 'Finished', not 'Active'), "
    "and (c) the data types are compatible with the operations (numeric columns for sums "
    "and comparisons).",
    space_after=Pt(6)
)

add_para("4. Encode as executable code:", bold=True, space_after=Pt(4))
add_para(
    "The verified logic is encoded in deterministic_engine.py as pure pandas operations. "
    "No LLM is involved. The code directly implements the three criteria using groupby, "
    "sum, nunique, and set operations. Every line of code can be traced back to a specific "
    "element in the Collibra governance catalog.",
    space_after=Pt(6)
)

add_para("5. Document provenance:", bold=True, space_after=Pt(4))
add_para(
    "The engine returns not just the answer (303) but a full provenance dictionary: which "
    "Collibra assets were used, which tables and columns were queried, how each criterion "
    "was computed, and how many participants met each criterion. This makes the answer "
    "auditable and defensible.",
    space_after=Pt(8)
)

doc.add_heading("The Dual Approach", level=2)
add_para(
    "Our solution uses both approaches in tandem:",
    space_after=Pt(4)
)
add_bullet("RAG Pipeline (LLM-based): For open-ended, exploratory questions where the user "
           "might ask anything. The LLM uses governed context to generate SQL on the fly. "
           "Flexible but has a small margin of variability.")
add_bullet("Deterministic Engine (rulebook-based): For known, governed metrics where the "
           "definition is fully specified. Produces identical results every time. No LLM, "
           "no variability, full auditability.")
add_para(
    "This dual approach gives us the best of both worlds: the flexibility of AI for new "
    "questions and the reliability of deterministic code for mission-critical metrics.",
    space_after=Pt(8)
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 7. RESULTS SUMMARY
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading("7. Results Summary", level=1)

table2 = doc.add_table(rows=8, cols=2, style="Light Shading Accent 1")
table2.rows[0].cells[0].text = "Metric"
table2.rows[0].cells[1].text = "Value"
results_data = [
    ["Active testers in ongoing programs", "303"],
    ["Criterion 1 (tickets >= 3)", "73 participants"],
    ["Criterion 2 (surveys >= 2)", "284 participants"],
    ["Criterion 3 (completed > rest)", "64 participants"],
    ["Determinism test runs", "20/20 identical"],
    ["Overlapping participants removed", "118"],
    ["Total unique participants in ongoing programs", "~1,800"],
]
for r_idx, (k, v) in enumerate(results_data):
    table2.rows[r_idx + 1].cells[0].text = k
    table2.rows[r_idx + 1].cells[1].text = v

doc.add_paragraph()

add_para("Data sources used:", bold=True)
add_bullet("zcc_prt_mtrc — Participant Metrics (5,865 rows total)")
add_bullet("zcc_tkt_itm — Feedback Tickets (3,852 rows total)")
add_bullet("zcc_prj_hdr — Project Headers (100 projects: 80 Finished, 20 Ongoing)")
add_bullet("Governed context from Collibra (52 business terms, 4 measures, 65 column mappings)")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 8. HOW TO REPLICATE THIS PROJECT
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading("8. How to Replicate This Project", level=1)

add_para(
    "Follow these steps in order. Each step builds on the previous one.",
    space_after=Pt(8)
)

steps = [
    ("Clone and set up", 
     "Clone the repository and create a virtual environment.\n"
     "  git clone https://github.com/erkantrt/NextLevelChallenge.git\n"
     "  cd NextLevelChallenge\n"
     "  python -m venv .venv\n"
     "  .venv\\Scripts\\activate  (Windows) or source .venv/bin/activate (macOS/Linux)\n"
     "  pip install delta-sharing pandas requests anthropic beautifulsoup4 lxml python-dotenv pandasql"),
    ("Configure credentials",
     "Create a .env file with your Collibra username/password and Anthropic API key. "
     "Place your Delta Sharing config.json in the project root."),
    ("Verify connections",
     "Test each API independently: list Delta Sharing tables, call Collibra /rest/2.0/communities, "
     "send a test prompt to Claude. All should return HTTP 200 or valid responses."),
    ("Ingest governed context",
     "Run: python collibra_client.py\n"
     "This creates governed_context.json with all business terms, measures, column mappings, "
     "and relations from Collibra."),
    ("Load and explore data",
     "Run: python data_loader.py\n"
     "This loads all 8 tables from Databricks and prints their schemas. Explore the data "
     "to understand what is available."),
    ("Run the before/after demo",
     "Run: python ask_your_data.py\n"
     "This demonstrates the LLM answering with and without governed context. Observe how "
     "the same question produces drastically different answers."),
    ("Run the deterministic engine",
     "Run: python test_deterministic.py\n"
     "This runs the Active Tester computation 20 times and verifies all runs return 303."),
    ("Extend",
     "To add new governed rules, follow the pattern in rulebook.py: extract the definition "
     "from Collibra, verify column names against actual data, encode the logic in Python, "
     "and document the provenance."),
]

for i, (title, desc) in enumerate(steps, 1):
    add_para(f"Step {i}: {title}", bold=True, space_after=Pt(4))
    add_para(desc, space_after=Pt(8))

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 9. APPENDIX: FILE REFERENCE
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading("9. Appendix: File Reference", level=1)

files = [
    ("collibra_client.py",
     "Connects to the Collibra REST API and extracts all governed context "
     "(business terms, measures, logical assets, physical assets, relations, column mappings). "
     "Saves to governed_context.json. Handles HTML stripping and pagination."),
    ("data_loader.py",
     "Loads all 8 Databricks tables via Delta Sharing into pandas DataFrames. "
     "Provides a SQL execution interface via pandasql. Caches loaded data in memory."),
    ("ask_your_data.py",
     "The main RAG pipeline. Builds LLM prompts with and without governed context, "
     "sends questions to Claude, extracts SQL from responses, executes it, and returns "
     "answers with full traceability. Used for the before/after demonstration."),
    ("rulebook.py",
     "Encodes Collibra's governed definitions as Python constants. Contains the "
     "verbatim Active Tester definition, Active Tester Flag criteria, project status "
     "values, and a comprehensive column map. Every entry includes the Collibra asset ID."),
    ("deterministic_engine.py",
     "Pure pandas computation of the Active Tester count. Implements the three criteria "
     "from the rulebook using groupby, sum, nunique, and set union. No LLM involved. "
     "Returns the answer plus full provenance."),
    ("test_deterministic.py",
     "Test harness that runs the deterministic engine N times and verifies all results "
     "are identical. Saves results to deterministic_test_results.json."),
    ("governed_context.json",
     "Cached output of collibra_client.py. Contains 52 business terms, 4 measures, "
     "52 logical assets, 74 physical assets, 230 relations, and 65 column mappings."),
    ("config.json",
     "Delta Sharing profile containing the bearer token and Databricks endpoint URL. "
     "Grants read-only access to 8 tables in centercode_share.centercode."),
    (".env",
     "Environment variables: Collibra credentials, Anthropic API key, Delta Sharing "
     "profile path, and row limit configuration."),
    ("discover.py / discover2.py",
     "One-time data discovery scripts used to verify column names, enumerate distinct "
     "values (e.g., Z_PRJ_STAT: 'Ongoing'/'Finished'), and trace Collibra relations "
     "for the Active Tester Flag."),
]

for fname, desc in files:
    add_para(fname, bold=True, space_after=Pt(2))
    add_para(desc, space_after=Pt(8))


# ═══════════════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════════════
output_path = os.path.join(os.path.dirname(__file__), "Collibra_Ask_Your_Data_Project_Guide.docx")
doc.save(output_path)
print(f"Document saved to: {output_path}")
